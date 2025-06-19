"""
Content processor service for extracting and processing content from various file types.

This module handles content extraction from PDFs, DOCX, text files, and other document formats,
providing structured data for search indexing and AI processing.
"""
import asyncio
import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Any, Optional

logger = logging.getLogger(__name__)


@dataclass
class ContentExtractionResult:
    """Result of content extraction from a file"""
    text: str
    summary: str
    sections: Dict[str, Any]
    metadata: Dict[str, Any]
    key_points: Optional[List[str]] = None
    tables: Optional[List[Dict[str, Any]]] = None
    confidence_score: float = 0.0


class ContentProcessor:
    """Service for extracting and processing content from various file types"""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.max_text_length = self.config.get('max_text_length', 1_000_000)  # 1MB of text
        
        # Initialize extractors
        self.extractors = {
            'application/pdf': self._extract_pdf,
            'application/msword': self._extract_doc,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx,
            'text/plain': self._extract_text,
            'text/markdown': self._extract_markdown,
            'text/html': self._extract_html,
            'text/csv': self._extract_csv,
            'application/json': self._extract_json,
        }
    
    async def extract_content(
        self, 
        file_content: bytes, 
        content_type: str,
        file_name: str = None
    ) -> ContentExtractionResult:
        """Extract content based on file type"""
        if content_type not in self.extractors:
            # Try to detect from filename
            if file_name:
                ext = Path(file_name).suffix.lower()
                content_type = self._get_content_type_from_extension(ext)
            
            if content_type not in self.extractors:
                raise ValueError(f"Unsupported content type: {content_type}")
        
        try:
            extractor = self.extractors[content_type]
            result = await extractor(file_content, file_name or 'unknown')
            
            # Post-process the results
            result = await self._post_process_content(result)
            
            return result
            
        except Exception as e:
            logger.error(f"Content extraction failed for {content_type}: {str(e)}")
            # Return basic result for fallback
            return ContentExtractionResult(
                text=f"Content extraction failed: {str(e)}",
                summary="Failed to extract content",
                sections={},
                metadata={"error": str(e), "content_type": content_type},
                confidence_score=0.0
            )
    
    async def _extract_pdf(self, content: bytes, file_name: str) -> ContentExtractionResult:
        """Extract content from PDF files"""
        try:
            from io import BytesIO
            
            # Try PyPDF2 first for basic extraction
            try:
                import PyPDF2
                
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                text_parts = []
                sections = {}
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                        sections[f"page_{i+1}"] = {
                            "content": page_text,
                            "page_number": i + 1,
                            "type": "page"
                        }
                
                full_text = "\n".join(text_parts)
                
                # Generate summary and metadata
                summary = await self._generate_summary(full_text)
                key_points = await self._extract_key_points(full_text)
                
                metadata = {
                    "pages": len(pdf_reader.pages),
                    "file_type": "PDF",
                    "estimated_reading_time": len(full_text.split()) // 200  # Assume 200 WPM
                }
                
                return ContentExtractionResult(
                    text=full_text,
                    summary=summary,
                    sections=sections,
                    metadata=metadata,
                    key_points=key_points,
                    confidence_score=0.8 if full_text else 0.2
                )
                
            except ImportError:
                logger.warning("PyPDF2 not available, using basic text extraction")
                # Fallback to basic text extraction
                text = content.decode('utf-8', errors='ignore')
                summary = await self._generate_summary(text)
                
                return ContentExtractionResult(
                    text=text,
                    summary=summary,
                    sections={"full_content": {"content": text, "type": "raw"}},
                    metadata={"file_type": "PDF", "extraction_method": "fallback"},
                    confidence_score=0.3
                )
                
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {str(e)}")
    
    async def _extract_docx(self, content: bytes, file_name: str) -> ContentExtractionResult:
        """Extract content from DOCX files"""
        try:
            import docx
            from io import BytesIO
            
            doc = docx.Document(BytesIO(content))
            text_parts = []
            sections = {}
            current_section = {"title": "Document Start", "content": "", "level": 0}
            section_counter = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
                    # Check if it's a heading
                    if paragraph.style.name.startswith('Heading'):
                        # Save previous section
                        if current_section["content"]:
                            sections[f"section_{section_counter}"] = current_section.copy()
                            section_counter += 1
                        
                        # Start new section
                        level = 1
                        try:
                            level = int(paragraph.style.name.replace('Heading ', ''))
                        except ValueError:
                            pass
                        
                        current_section = {
                            "title": paragraph.text,
                            "content": "",
                            "level": level,
                            "type": "heading"
                        }
                    else:
                        current_section["content"] += paragraph.text + "\n"
            
            # Save final section
            if current_section["content"]:
                sections[f"section_{section_counter}"] = current_section
            
            full_text = "\n".join(text_parts)
            summary = await self._generate_summary(full_text)
            key_points = await self._extract_key_points(full_text)
            
            metadata = {
                "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
                "sections": len(sections),
                "file_type": "DOCX"
            }
            
            return ContentExtractionResult(
                text=full_text,
                summary=summary,
                sections=sections,
                metadata=metadata,
                key_points=key_points,
                confidence_score=0.9
            )
            
        except ImportError:
            logger.warning("python-docx not available")
            raise ValueError("DOCX extraction requires python-docx library")
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {str(e)}")
    
    async def _extract_text(self, content: bytes, file_name: str) -> ContentExtractionResult:
        """Extract content from plain text files"""
        try:
            text = content.decode('utf-8')
            
            # Simple section detection based on multiple newlines
            sections = {}
            paragraphs = text.split('\n\n')
            
            for i, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    sections[f"paragraph_{i}"] = {
                        "title": f"Paragraph {i+1}",
                        "content": paragraph.strip(),
                        "type": "paragraph"
                    }
            
            summary = await self._generate_summary(text)
            key_points = await self._extract_key_points(text)
            
            metadata = {
                "character_count": len(text),
                "word_count": len(text.split()),
                "paragraphs": len(paragraphs),
                "file_type": "Plain Text"
            }
            
            return ContentExtractionResult(
                text=text,
                summary=summary,
                sections=sections,
                metadata=metadata,
                key_points=key_points,
                confidence_score=1.0
            )
            
        except Exception as e:
            raise ValueError(f"Text extraction failed: {str(e)}")
    
    async def _extract_markdown(self, content: bytes, file_name: str) -> ContentExtractionResult:
        """Extract content from Markdown files"""
        try:
            text = content.decode('utf-8')
            
            # Extract sections based on headers
            sections = {}
            current_section = {"title": "Introduction", "content": "", "level": 0}
            section_counter = 0
            
            lines = text.split('\n')
            for line in lines:
                header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
                if header_match:
                    # Save previous section
                    if current_section["content"].strip():
                        sections[f"section_{section_counter}"] = current_section.copy()
                        section_counter += 1
                    
                    # Start new section
                    level = len(header_match.group(1))
                    title = header_match.group(2)
                    current_section = {
                        "title": title,
                        "content": "",
                        "level": level,
                        "type": "markdown_header"
                    }
                else:
                    current_section["content"] += line + "\n"
            
            # Save final section
            if current_section["content"].strip():
                sections[f"section_{section_counter}"] = current_section
            
            summary = await self._generate_summary(text)
            key_points = await self._extract_key_points(text)
            
            metadata = {
                "sections": len(sections),
                "file_type": "Markdown"
            }
            
            return ContentExtractionResult(
                text=text,
                summary=summary,
                sections=sections,
                metadata=metadata,
                key_points=key_points,
                confidence_score=0.95
            )
            
        except Exception as e:
            raise ValueError(f"Markdown extraction failed: {str(e)}")
    
    async def _extract_html(self, content: bytes, file_name: str) -> ContentExtractionResult:
        """Extract content from HTML files"""
        try:
            text = content.decode('utf-8')
            
            # Basic HTML tag removal (simplified)
            import re
            # Remove script and style elements
            text = re.sub(r'<script.*?</script>', '', text, flags=re.DOTALL)
            text = re.sub(r'<style.*?</style>', '', text, flags=re.DOTALL)
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', '', text)
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            summary = await self._generate_summary(text)
            
            sections = {
                "html_content": {
                    "title": f"HTML Content - {file_name}",
                    "content": text,
                    "type": "html_document"
                }
            }
            
            metadata = {
                "file_type": "HTML",
                "word_count": len(text.split())
            }
            
            return ContentExtractionResult(
                text=text,
                summary=summary,
                sections=sections,
                metadata=metadata,
                confidence_score=0.7  # Lower confidence due to basic extraction
            )
            
        except Exception as e:
            raise ValueError(f"HTML extraction failed: {str(e)}")
    
    async def _extract_csv(self, content: bytes, file_name: str) -> ContentExtractionResult:
        """Extract content from CSV files"""
        try:
            import csv
            from io import StringIO
            
            text = content.decode('utf-8')
            csv_reader = csv.reader(StringIO(text))
            
            rows = list(csv_reader)
            if not rows:
                raise ValueError("Empty CSV file")
            
            # Convert to text representation
            text_representation = "\n".join(["\t".join(row) for row in rows])
            
            # Create structured representation
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []
            
            sections = {
                "csv_data": {
                    "title": f"CSV Data - {file_name}",
                    "content": text_representation,
                    "type": "csv_table",
                    "headers": headers,
                    "data_rows": len(data_rows)
                }
            }
            
            summary = f"CSV file with {len(headers)} columns and {len(data_rows)} data rows"
            
            metadata = {
                "rows": len(rows),
                "columns": len(headers),
                "has_headers": True,
                "file_type": "CSV"
            }
            
            return ContentExtractionResult(
                text=text_representation,
                summary=summary,
                sections=sections,
                metadata=metadata,
                confidence_score=0.95
            )
            
        except Exception as e:
            raise ValueError(f"CSV extraction failed: {str(e)}")
    
    async def _extract_json(self, content: bytes, file_name: str) -> ContentExtractionResult:
        """Extract content from JSON files"""
        try:
            import json
            
            text = content.decode('utf-8')
            data = json.loads(text)
            
            # Convert JSON to readable text
            text_representation = json.dumps(data, indent=2, ensure_ascii=False)
            
            # Create summary based on structure
            def analyze_json_structure(obj):
                if isinstance(obj, dict):
                    return f"Object with {len(obj)} keys"
                elif isinstance(obj, list):
                    return f"Array with {len(obj)} items"
                else:
                    return f"Value: {str(obj)[:50]}"
            
            summary = f"JSON data: {analyze_json_structure(data)}"
            
            sections = {
                "json_structure": {
                    "title": f"JSON Structure - {file_name}",
                    "content": text_representation,
                    "type": "json_data"
                }
            }
            
            metadata = {
                "json_type": type(data).__name__,
                "file_type": "JSON"
            }
            
            if isinstance(data, dict):
                metadata["keys"] = len(data)
            elif isinstance(data, list):
                metadata["items"] = len(data)
            
            return ContentExtractionResult(
                text=text_representation,
                summary=summary,
                sections=sections,
                metadata=metadata,
                confidence_score=1.0
            )
            
        except Exception as e:
            raise ValueError(f"JSON extraction failed: {str(e)}")
    
    async def _extract_doc(self, content: bytes, file_name: str) -> ContentExtractionResult:
        """Extract content from legacy DOC files"""
        # For now, we'll return a placeholder
        # In production, you might use python-docx2txt or antiword
        text = "Legacy DOC format - content extraction requires additional dependencies"
        
        sections = {
            "doc_placeholder": {
                "title": f"DOC File - {file_name}",
                "content": text,
                "type": "doc_legacy"
            }
        }
        
        metadata = {
            "file_type": "DOC (Legacy)",
            "extraction_note": "Requires additional libraries for full extraction"
        }
        
        return ContentExtractionResult(
            text=text,
            summary="Legacy DOC file detected",
            sections=sections,
            metadata=metadata,
            confidence_score=0.1
        )
    
    async def _post_process_content(self, result: ContentExtractionResult) -> ContentExtractionResult:
        """Post-process extracted content for quality and consistency"""
        
        # Truncate text if too long
        if len(result.text) > self.max_text_length:
            result.text = result.text[:self.max_text_length] + "... [Content truncated]"
            result.metadata["truncated"] = True
        
        # Ensure summary is not too long
        if len(result.summary) > 1000:
            result.summary = result.summary[:1000] + "..."
        
        return result
    
    async def _generate_summary(self, text: str, max_length: int = 500) -> str:
        """Generate a summary of the text (basic implementation)"""
        if not text.strip():
            return "Empty content"
        
        # Basic extractive summarization
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        if not sentences:
            return text[:max_length] + "..." if len(text) > max_length else text
        
        # Take first few sentences up to max_length
        summary_parts = []
        current_length = 0
        
        for sentence in sentences[:10]:  # Limit to first 10 sentences
            if current_length + len(sentence) < max_length:
                summary_parts.append(sentence)
                current_length += len(sentence)
            else:
                break
        
        summary = ". ".join(summary_parts)
        if summary and not summary.endswith('.'):
            summary += "."
        
        return summary or text[:max_length] + "..."
    
    async def _extract_key_points(self, text: str, max_points: int = 10) -> List[str]:
        """Extract key points from text (basic implementation)"""
        if not text.strip():
            return []
        
        # Simple approach: find sentences with key phrases
        key_phrases = ['important', 'significant', 'key', 'critical', 'essential', 'main', 'primary']
        
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        key_points = []
        
        for sentence in sentences:
            if any(phrase in sentence.lower() for phrase in key_phrases):
                if len(sentence) > 20 and len(sentence) < 200:  # Reasonable length
                    key_points.append(sentence + ".")
                    if len(key_points) >= max_points:
                        break
        
        # If no key phrases found, take first few substantive sentences
        if not key_points:
            key_points = [s + "." for s in sentences[:3] if len(s) > 20]
        
        return key_points
    
    def _get_content_type_from_extension(self, ext: str) -> str:
        """Get content type from file extension"""
        extension_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.html': 'text/html',
            '.htm': 'text/html',
            '.csv': 'text/csv',
            '.json': 'application/json',
        }
        return extension_map.get(ext.lower(), 'application/octet-stream')
